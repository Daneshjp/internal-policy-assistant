"""Document service for CRUD operations and text extraction."""

import logging
from pathlib import Path
from typing import Optional, List

from fastapi import UploadFile
from sqlalchemy import or_, func
from sqlalchemy.orm import Session, joinedload

from app.models.document import Document, DocumentStatus
from app.models.category import Category
from app.schemas.document import DocumentCreate, DocumentUpdate, DocumentFilters
from app.services.storage_service import storage_service
from app.exceptions import NotFoundError, ValidationError

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for document-related operations."""

    def get_documents(
        self,
        db: Session,
        user_id: Optional[int] = None,
        filters: Optional[DocumentFilters] = None,
        page: int = 1,
        per_page: int = 20
    ) -> tuple[List[Document], int]:
        """
        Get paginated list of documents with optional filters.

        Args:
            db: Database session
            user_id: Optional user ID to filter by owner
            filters: Optional filter criteria
            page: Page number (1-indexed)
            per_page: Number of items per page

        Returns:
            Tuple of (documents, total_count)
        """
        query = db.query(Document).options(
            joinedload(Document.user),
            joinedload(Document.category)
        )

        # Apply user filter if provided
        if user_id is not None:
            query = query.filter(Document.user_id == user_id)

        # Apply filters if provided
        if filters:
            if filters.category_id is not None:
                query = query.filter(Document.category_id == filters.category_id)

            if filters.status is not None:
                query = query.filter(Document.status == filters.status)

            if filters.search:
                search_term = f"%{filters.search}%"
                query = query.filter(
                    or_(
                        Document.title.ilike(search_term),
                        Document.description.ilike(search_term)
                    )
                )

        # Get total count before pagination
        total = query.count()

        # Apply pagination
        offset = (page - 1) * per_page
        documents = query.order_by(Document.created_at.desc()).offset(offset).limit(per_page).all()

        return documents, total

    def get_document(self, db: Session, document_id: int) -> Document:
        """
        Get a single document by ID.

        Args:
            db: Database session
            document_id: Document ID

        Returns:
            Document instance

        Raises:
            NotFoundError: If document not found
        """
        document = db.query(Document).options(
            joinedload(Document.user),
            joinedload(Document.category)
        ).filter(Document.id == document_id).first()

        if not document:
            raise NotFoundError("Document")

        return document

    async def create_document(
        self,
        db: Session,
        user_id: int,
        data: DocumentCreate,
        file: UploadFile
    ) -> Document:
        """
        Create a new document with file upload.

        Args:
            db: Database session
            user_id: ID of the uploading user
            data: Document creation data
            file: Uploaded file

        Returns:
            Created document instance
        """
        # Validate category if provided
        if data.category_id is not None:
            category = db.query(Category).filter(Category.id == data.category_id).first()
            if not category:
                raise ValidationError(f"Category with ID {data.category_id} not found")

        # Save file to storage
        file_path, file_size = await storage_service.save_file(file, user_id)
        file_type = storage_service.get_file_type(file.filename)

        # Create document record
        document = Document(
            user_id=user_id,
            category_id=data.category_id,
            title=data.title,
            description=data.description,
            file_url=file_path,
            file_type=file_type,
            file_size=file_size,
            status=DocumentStatus.processing
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        # Extract text content asynchronously (in background)
        # For now, we'll do it synchronously
        try:
            full_path = storage_service.get_file_path(file_path)
            if full_path:
                content = self.extract_text_from_file(str(full_path), file_type)
                document.content = content
                document.status = DocumentStatus.active
                db.commit()
                db.refresh(document)
        except Exception as e:
            logger.error(f"Failed to extract text from document {document.id}: {e}")
            document.status = DocumentStatus.active  # Still mark as active
            db.commit()

        # Reload with relationships
        return self.get_document(db, document.id)

    def update_document(
        self,
        db: Session,
        document_id: int,
        data: DocumentUpdate
    ) -> Document:
        """
        Update a document.

        Args:
            db: Database session
            document_id: Document ID
            data: Update data

        Returns:
            Updated document instance

        Raises:
            NotFoundError: If document not found
        """
        document = db.query(Document).filter(Document.id == document_id).first()

        if not document:
            raise NotFoundError("Document")

        # Validate category if provided
        if data.category_id is not None:
            category = db.query(Category).filter(Category.id == data.category_id).first()
            if not category:
                raise ValidationError(f"Category with ID {data.category_id} not found")

        # Update fields if provided
        update_data = data.model_dump(exclude_unset=True)
        for field, value in update_data.items():
            setattr(document, field, value)

        db.commit()
        db.refresh(document)

        return self.get_document(db, document.id)

    def delete_document(self, db: Session, document_id: int) -> bool:
        """
        Delete a document and its associated file.

        Args:
            db: Database session
            document_id: Document ID

        Returns:
            True if deleted successfully

        Raises:
            NotFoundError: If document not found
        """
        document = db.query(Document).filter(Document.id == document_id).first()

        if not document:
            raise NotFoundError("Document")

        # Delete file from storage
        if document.file_url:
            storage_service.delete_file(document.file_url)

        # Delete document record
        db.delete(document)
        db.commit()

        return True

    def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """
        Extract text content from a file.

        Args:
            file_path: Full path to the file
            file_type: Type of file (pdf, docx, txt)

        Returns:
            Extracted text content
        """
        file_type = file_type.lower()

        if file_type == "pdf":
            return self._extract_text_from_pdf(file_path)
        elif file_type == "docx":
            return self._extract_text_from_docx(file_path)
        elif file_type == "txt":
            return self._extract_text_from_txt(file_path)
        else:
            logger.warning(f"Unsupported file type for text extraction: {file_type}")
            return ""

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using PyPDF2."""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)
        except ImportError:
            logger.error("PyPDF2 not installed. Install it with: pip install PyPDF2")
            return ""
        except Exception as e:
            logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            return ""

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file using python-docx."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            return "\n\n".join(text_parts)
        except ImportError:
            logger.error("python-docx not installed. Install it with: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            return ""

    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Failed to read text file {file_path}: {e}")
            return ""


# Singleton instance
document_service = DocumentService()
